import csv
import io
import json
import logging
import re
import fitz  # PyMuPDF
import tiktoken
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from langsmith import traceable
from app.config import get_settings
from app.database import get_supabase_client
from app.services.embedding_service import EmbeddingService
from app.services.metadata_service import MetadataService

logger = logging.getLogger(__name__)

settings = get_settings()
embedding_service = EmbeddingService()
metadata_service = MetadataService()


def parse_text(file_bytes: bytes, mime_type: str) -> str:
    if mime_type == "application/pdf":
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = "".join(page.get_text() for page in doc)
        doc.close()
        return text
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _parse_docx(file_bytes)
    if mime_type == "text/csv":
        return _parse_csv(file_bytes)
    if mime_type == "text/html":
        return _parse_html(file_bytes)
    if mime_type == "application/json":
        return _parse_json(file_bytes)
    # TXT or Markdown
    return file_bytes.decode("utf-8", errors="ignore")


def _parse_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_csv(file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8", errors="ignore")
    reader = csv.reader(io.StringIO(text))
    return "\n".join(" | ".join(row) for row in reader)


def _parse_html(file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8", errors="ignore")
    soup = BeautifulSoup(text, "html.parser")
    for element in soup(["script", "style"]):
        element.decompose()
    return soup.get_text(separator="\n", strip=True)


def _parse_json(file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8", errors="ignore")
    data = json.loads(text)
    return json.dumps(data, indent=2, ensure_ascii=False)


def chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    """Split text respecting document structure boundaries.

    Tries structural splits first (BAB, Pasal, numbered sections, paragraphs).
    Falls back to sliding window for sections that are still too large.
    """
    enc = tiktoken.get_encoding("cl100k_base")

    # Structural boundary patterns (most specific first)
    structure_patterns = [
        r'\n(?=BAB\s+[IVXLCDM]+)',       # Chapter: BAB I, BAB II
        r'\n(?=Pasal\s+\d+)',              # Article: Pasal 1, Pasal 2
        r'\n(?=PASAL\s+\d+)',              # Uppercase variant
        r'\n(?=Bagian\s+\w+)',             # Section: Bagian Kesatu
        r'\n(?=\d+\.\s+[A-Z])',            # Numbered section: 1. Title
        r'\n\n',                            # Double newline (paragraph)
    ]

    sections = [text]
    for pattern in structure_patterns:
        new_sections = []
        for section in sections:
            parts = re.split(pattern, section)
            new_sections.extend(p for p in parts if p.strip())
        # Accept this split level if every section fits within 1.5x chunk_size
        if len(new_sections) > 1 and all(
            len(enc.encode(s)) <= chunk_size * 1.5 for s in new_sections
        ):
            sections = new_sections
            break

    # For sections still too large, fall back to sliding window
    chunks = []
    for section in sections:
        tokens = enc.encode(section)
        if len(tokens) <= chunk_size:
            chunks.append(section)
        else:
            start = 0
            while start < len(tokens):
                end = min(start + chunk_size, len(tokens))
                chunk = enc.decode(tokens[start:end])
                if chunk.strip():
                    chunks.append(chunk)
                if end == len(tokens):
                    break
                start += chunk_size - chunk_overlap

    return [c for c in chunks if c.strip()]


def _contextualize_chunks(chunks: list[str], metadata: dict | None) -> list[str]:
    """Prepend document context to each chunk for richer embeddings.

    The original chunk text is stored in the DB; the contextualized version
    is used only for embedding generation.
    """
    if not metadata:
        return chunks

    parts = []
    if metadata.get("title"):
        parts.append(f"Document: {metadata['title']}")
    if metadata.get("category"):
        parts.append(f"Category: {metadata['category']}")
    if metadata.get("author"):
        parts.append(f"Author: {metadata['author']}")

    if not parts:
        return chunks

    header = "[" + " | ".join(parts) + "]"
    return [f"{header}\n\n{chunk}" for chunk in chunks]


@traceable
async def process_document(
    doc_id: str,
    user_id: str,
    file_path: str,
    mime_type: str,
    embedding_model: str | None = None,
    llm_model: str | None = None,
):
    client = get_supabase_client()
    try:
        client.table("documents").update({"status": "processing"}).eq("id", doc_id).eq("user_id", user_id).execute()

        # Download from Supabase Storage
        file_bytes = client.storage.from_(settings.storage_bucket).download(file_path)

        # Parse text
        text = parse_text(file_bytes, mime_type)
        if not text.strip():
            raise ValueError("No text content extracted from file")

        # Extract metadata (best-effort — failures are logged but don't block ingestion)
        metadata_dict = None
        try:
            metadata = await metadata_service.extract_metadata(text, model=llm_model)
            if metadata:
                metadata_dict = metadata.model_dump()
        except Exception as e:
            logger.warning("Metadata extraction skipped for doc_id=%s: %s", doc_id, e)

        # Chunk and embed
        chunks = chunk_text(text, settings.rag_chunk_size, settings.rag_chunk_overlap)
        if not chunks:
            raise ValueError("No chunks generated")

        # Contextualize chunks for richer embeddings (original text stored in DB)
        embed_texts = _contextualize_chunks(chunks, metadata_dict)

        # Embed in batches of 100 using the user's chosen embedding model
        all_embeddings: list[list[float]] = []
        for i in range(0, len(embed_texts), 100):
            batch_embeddings = await embedding_service.embed_batch(
                embed_texts[i : i + 100], model=embedding_model
            )
            all_embeddings.extend(batch_embeddings)

        # Bulk insert chunks
        client.table("document_chunks").insert([
            {
                "document_id": doc_id,
                "user_id": user_id,
                "content": chunk,
                "chunk_index": i,
                "embedding": embedding,
            }
            for i, (chunk, embedding) in enumerate(zip(chunks, all_embeddings))
        ]).execute()

        # GraphRAG entity extraction (best-effort — failures logged, don't block)
        try:
            from app.services.system_settings_service import get_system_settings
            sys_settings = get_system_settings()
            if sys_settings.get("graph_enabled"):
                from app.services.graph_service import GraphService
                from app.services.audit_service import log_action
                graph_service = GraphService()

                inserted = client.table("document_chunks") \
                    .select("id") \
                    .eq("document_id", doc_id) \
                    .eq("user_id", user_id) \
                    .order("chunk_index") \
                    .execute()
                chunk_ids = [row["id"] for row in inserted.data or []]

                graph_model = sys_settings.get("graph_entity_extraction_model") or llm_model
                extraction = await graph_service.extract_entities(
                    chunks=chunks, doc_metadata=metadata_dict, model=graph_model,
                )
                if extraction.entities:
                    await graph_service.store_entities(
                        extraction=extraction, doc_id=doc_id, user_id=user_id,
                        chunk_ids=chunk_ids, chunks=chunks,
                    )
                    log_action(user_id, None, "graph_entities_extracted", "document", doc_id)
                    logger.info(
                        "Graph: %d entities, %d relationships for doc_id=%s",
                        len(extraction.entities), len(extraction.relationships), doc_id,
                    )
        except Exception as e:
            logger.warning("Graph entity extraction skipped for doc_id=%s: %s", doc_id, e)

        client.table("documents").update({
            "status": "completed",
            "chunk_count": len(chunks),
            "metadata": metadata_dict,
        }).eq("id", doc_id).eq("user_id", user_id).execute()

    except Exception as e:
        logger.error("Document processing failed for doc_id=%s: %s", doc_id, e)
        client.table("documents").update({
            "status": "failed",
            "error_msg": "Processing failed. Please try uploading again.",
        }).eq("id", doc_id).eq("user_id", user_id).execute()
