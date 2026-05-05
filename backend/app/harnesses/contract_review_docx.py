"""Phase 22 / DOCX-01..08 — Sandbox-driven .docx generation post_execute callback.

Called by harness_engine after CR-08 (executive-summary) completes its LLM_SINGLE call.

REVIEW #6: this callback runs `_render_summary_markdown` first to produce the
deterministic human-readable contract-review-report.md, THEN proceeds to DOCX.
Without this rendering, `contract-review-report.md` would never be written (CR-08's
LLM_SINGLE writes JSON to executive-summary.json, not markdown).

REVIEW #7: returns wrote_binary=True + size_bytes so the engine emits a
workspace_updated event after the DOCX write — the frontend Workspace Panel
auto-refreshes.

D-22-15: NEVER raise — error dict + fallback_message on any failure.

ISSUE-05 PIN: SandboxService.execute(*, code, thread_id, user_id, token).
  Returns {stdout, stderr, exit_code, error_type, execution_ms, files, execution_id}
  files[i] = {filename, size_bytes, signed_url, storage_path}
  NO files= kwarg, NO timeout_seconds= kwarg, NO base64 extraction patterns.
"""
from __future__ import annotations

import json
import logging
from typing import Any

import httpx

from app.services import audit_service

logger = logging.getLogger(__name__)


class DocxGenerationError(RuntimeError):
    """Internal — caught in _generate_docx_post_execute and translated to D-22-15 fallback."""


# ---------------------------------------------------------------------------
# DOCX_GENERATION_SCRIPT_BODY — Python script that runs INSIDE the sandbox.
# Receives INPUT_DATA dict prepended at runtime. Writes to /sandbox/output/.
# Uses python-docx 1.1.2 (installed in sandbox Dockerfile via plan 22-01).
# D-22-13: pastel hex codes — E6F4EA (green), FEF7E0 (yellow), FCE8E6 (red).
# D-22-12: programmatic generation, no template.
# ---------------------------------------------------------------------------

DOCX_GENERATION_SCRIPT_BODY = r'''
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

data = INPUT_DATA  # noqa: F821 — defined by inputs_literal prefix at runtime
classification = data.get('classification') or {}
review_context_text = data.get('review_context') or ''
playbook = data.get('playbook') or {}
risks = data.get('risks') or []
redlines = data.get('redlines') or []
exec_summary = data.get('executive_summary') or {}

# Pastel risk colors — D-22-13
COLOR_GREEN_FILL  = 'E6F4EA'
COLOR_YELLOW_FILL = 'FEF7E0'
COLOR_RED_FILL    = 'FCE8E6'

doc = Document()

# Title page (DOCX-02)
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('CONFIDENTIAL')
title_run.bold = True
title_run.font.size = Pt(20)
title_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run('Contract Review Report').bold = True

doc.add_paragraph(f"Contract type: {classification.get('contract_type', 'Unknown')}")
doc.add_paragraph(f"Parties: {', '.join(classification.get('parties') or [])}")
doc.add_paragraph(f"Governing law: {classification.get('governing_law', '?')}")
doc.add_paragraph(f"Overall risk: {exec_summary.get('overall_risk', '?')}")
doc.add_page_break()

# Executive summary (DOCX-03)
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(exec_summary.get('recommendation', ''))
breakdown = exec_summary.get('risk_breakdown') or {}
doc.add_paragraph(
    f"Risk breakdown — RED: {breakdown.get('RED', 0)}, "
    f"YELLOW: {breakdown.get('YELLOW', 0)}, GREEN: {breakdown.get('GREEN', 0)}"
)

# Numbered key findings (DOCX-04)
doc.add_heading('Key Findings', level=1)
for i, finding in enumerate(exec_summary.get('key_findings') or [], 1):
    doc.add_paragraph(f"{i}. {finding}")

# Color-coded redline table (DOCX-05)
doc.add_heading('Detailed Redline Analysis', level=1)
if risks:
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Risk'
    hdr[1].text = 'Clause'
    hdr[2].text = 'Original'
    hdr[3].text = 'Proposed / Rationale'

    redline_by_idx = {r.get('clause_index'): r for r in (redlines or [])}

    for risk in risks:
        row = table.add_row().cells
        row[0].text = risk.get('risk_grade', '')
        row[1].text = f"{risk.get('clause_category', '')}: {risk.get('clause_heading', '')}"
        rl = redline_by_idx.get(risk.get('clause_index'))
        if rl:
            row[2].text = (rl.get('original_text') or '')[:1500]
            row[3].text = (rl.get('proposed_text') or '') + '\n\n' + (rl.get('rationale') or '')
        else:
            row[2].text = '(no redline — see rationale)'
            row[3].text = risk.get('rationale', '')

        grade = risk.get('risk_grade', '')
        if grade == 'GREEN':
            fill_color = COLOR_GREEN_FILL
        elif grade == 'YELLOW':
            fill_color = COLOR_YELLOW_FILL
        elif grade == 'RED':
            fill_color = COLOR_RED_FILL
        else:
            fill_color = None
        if fill_color:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), fill_color)
            row[0]._tc.get_or_add_tcPr().append(shd)
else:
    doc.add_paragraph('No risk-graded clauses available.')

# GREEN clauses (DOCX-06)
doc.add_heading('Acceptable Clauses (GREEN)', level=1)
greens = [r for r in risks if r.get('risk_grade') == 'GREEN']
if greens:
    for r in greens:
        doc.add_paragraph(
            f"- {r.get('clause_category', '')}: {r.get('clause_heading', '')} — no changes recommended.",
            style='List Bullet',
        )
else:
    doc.add_paragraph('None — every clause has at least YELLOW status.')

# Recommended next steps (DOCX-07)
doc.add_heading('Recommended Next Steps', level=1)
for i, step in enumerate(exec_summary.get('next_steps') or [], 1):
    doc.add_paragraph(f"{i}. {step}")

out_path = '/sandbox/output/contract-review.docx'
os.makedirs(os.path.dirname(out_path), exist_ok=True)
doc.save(out_path)
'''


# ---------------------------------------------------------------------------
# _generate_docx_post_execute — the post_execute callable
# ---------------------------------------------------------------------------

async def _generate_docx_post_execute(
    *,
    harness_run_id: str,
    thread_id: str,
    user_id: str,
    user_email: str,
    token: str,
    phase_results: dict,
    workspace,
    harness_name: str = "contract-review",
    **_,  # forward-compat for any future engine kwargs
) -> dict:
    """REVIEW #6 + #7 + D-22-15. NEVER raises.

    Steps:
    1. Read 6 workspace artifacts.
    2. REVIEW #6: call _render_summary_markdown to produce contract-review-report.md
       as real markdown (NOT raw JSON — LLM_SINGLE writes JSON to executive-summary.json).
    3. Build sandbox DOCX script + run via SandboxService.execute() (ISSUE-05 PIN).
    4. HTTP-GET the signed_url from sandbox output files.
    5. Write DOCX bytes to workspace via write_binary_file (source='harness').
    6. REVIEW #7: return wrote_binary=True + size_bytes so engine emits workspace_updated.
    7. D-22-15: any failure returns error dict; never raises.
    """
    from app.harnesses.contract_review import _render_summary_markdown

    run_id_short = (harness_run_id or "")[:8] or "unknown"
    out_path = f"contract-review-{run_id_short}.docx"

    try:
        # --- Step 1: Read all artifacts ---
        classification_md = (await workspace.read_file(thread_id, "classification.md")).get("content", "")
        review_ctx = (await workspace.read_file(thread_id, "review-context.md")).get("content", "")
        playbook_md = (await workspace.read_file(thread_id, "playbook-context.md")).get("content", "")
        risks_raw = (await workspace.read_file(thread_id, "risk-analysis.json")).get("content", "")
        redlines_raw = (await workspace.read_file(thread_id, "redlines.json")).get("content", "")
        exec_summary_raw = (await workspace.read_file(thread_id, "executive-summary.json")).get("content", "")

        classification = _extract_json_from_markdown(classification_md)
        playbook = _extract_json_from_markdown(playbook_md)
        risks = _extract_json_from_markdown(risks_raw, default=[])
        redlines = _extract_json_from_markdown(redlines_raw, default=[])
        executive_summary = _extract_json_from_markdown(exec_summary_raw)

        # --- Step 2 (REVIEW #6): render the deterministic markdown report BEFORE DOCX ---
        await _render_summary_markdown(
            executive_summary=executive_summary if isinstance(executive_summary, dict) else {},
            classification=classification if isinstance(classification, dict) else {},
            playbook=playbook if isinstance(playbook, dict) else {},
            risks=risks if isinstance(risks, list) else [],
            redlines=redlines if isinstance(redlines, list) else [],
            workspace=workspace,
            thread_id=thread_id,
        )

        # --- Step 3: Build sandbox inputs + run via ISSUE-05 pinned API ---
        sandbox_inputs = {
            "classification": classification if isinstance(classification, dict) else {},
            "review_context": review_ctx,
            "playbook": playbook if isinstance(playbook, dict) else {},
            "risks": risks if isinstance(risks, list) else [],
            "redlines": redlines if isinstance(redlines, list) else [],
            "executive_summary": executive_summary if isinstance(executive_summary, dict) else {},
        }

        from app.services.sandbox_service import get_sandbox_service
        sb = get_sandbox_service()

        # Prepend INPUT_DATA literal so the script body can reference it
        inputs_literal = f"INPUT_DATA = {json.dumps(sandbox_inputs, ensure_ascii=False)}\n"
        code_with_inputs = inputs_literal + DOCX_GENERATION_SCRIPT_BODY

        # ISSUE-05 PIN: execute(*, code, thread_id, user_id, token) — NO files=, NO timeout_seconds=
        sb_result = await sb.execute(
            code=code_with_inputs,
            thread_id=thread_id,
            user_id=user_id,
            token=token,
        )

        if sb_result.get("exit_code") != 0:
            stderr = (sb_result.get("stderr") or "")[:500]
            raise DocxGenerationError(
                f"sandbox exit_code={sb_result.get('exit_code')}: {stderr}"
            )

        sb_files = sb_result.get("files") or []
        if not sb_files:
            raise DocxGenerationError("sandbox produced no output files")

        docx_meta = next(
            (f for f in sb_files if f.get("filename") == "contract-review.docx"),
            None,
        )
        if docx_meta is None:
            raise DocxGenerationError(
                f"contract-review.docx missing in sandbox files: "
                f"{[f.get('filename') for f in sb_files]}"
            )
        sandbox_signed_url = docx_meta.get("signed_url") or ""
        if not sandbox_signed_url:
            raise DocxGenerationError("sandbox file missing signed_url")

        # --- Step 4: HTTP-GET the DOCX bytes from sandbox signed URL ---
        async with httpx.AsyncClient(timeout=60) as hc:
            resp = await hc.get(sandbox_signed_url)
            resp.raise_for_status()
            docx_bytes = resp.content
        if not docx_bytes:
            raise DocxGenerationError("sandbox returned empty DOCX bytes")

        # --- Step 5: Write to workspace (source='harness', D-22-14) ---
        await workspace.write_binary_file(
            thread_id=thread_id,
            file_path=out_path,
            content_bytes=docx_bytes,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            user_id=user_id,
            source="harness",
        )

        # --- Step 5b: get signed URL for chat chip (non-fatal if unavailable) ---
        signed_url = ""
        if hasattr(workspace, "get_signed_url"):
            try:
                signed_url = await workspace.get_signed_url(thread_id, out_path) or ""
            except Exception as exc:
                logger.warning("get_signed_url failed (non-fatal): %s", exc)

        # --- Step 6: audit success ---
        try:
            audit_service.log_action(
                user_id=user_id,
                user_email=user_email,
                action="contract_review_docx_generated",
                resource_type="harness_runs",
                resource_id=harness_run_id,
            )
        except Exception as exc:
            logger.warning("audit log failed (non-fatal): %s", exc)

        # REVIEW #7: wrote_binary=True signals plan 22-03 engine to emit workspace_updated
        return {
            "ok": True,
            "docx_path": out_path,
            "signed_url": signed_url,
            "wrote_binary": True,
            "size_bytes": len(docx_bytes),
        }

    except Exception as exc:
        # D-22-15: NEVER raise — return error dict; harness_runs.status stays 'completed'
        logger.warning(
            "docx post_execute failed harness_run=%s: %s",
            harness_run_id, exc, exc_info=True,
        )
        try:
            audit_service.log_action(
                user_id=user_id,
                user_email=user_email,
                action="contract_review_docx_failed",
                resource_type="harness_runs",
                resource_id=harness_run_id,
            )
        except Exception:
            pass
        return {
            "error": "docx_generation_failed",
            "code": "DOCX_FAILED",
            "detail": str(exc)[:500],
            "fallback_message": (
                "DOCX export unavailable right now — the full markdown summary is above. "
                "Retry by re-running the harness if needed."
            ),
        }


# ---------------------------------------------------------------------------
# _extract_json_from_markdown — shared parse helper
# ---------------------------------------------------------------------------

def _extract_json_from_markdown(md_text: str, default: Any = None) -> Any:
    """Extract JSON from markdown text — tries fenced block, then direct parse.

    Used to extract JSON from workspace files that may be wrapped in ```json``` fences
    (e.g., LLM_SINGLE output, LLM_AGENT markdown output).
    """
    if default is None:
        default = {}
    if not md_text:
        return default
    import re as _re
    m = _re.search(r"```json\s*(\{.*?\}|\[.*?\])\s*```", md_text, _re.DOTALL)
    if m:
        try:
            return json.loads(m.group(1))
        except Exception:
            pass
    try:
        return json.loads(md_text)
    except Exception:
        pass
    return default
