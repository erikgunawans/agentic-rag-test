"""Generate backend/tests/data/synth-contract.docx — 3-clause MSA fixture for E2E tests.

Run this script manually to regenerate the fixture:
    cd backend && python tests/data/_generate_synth_contract.py

The fixture is pre-committed to the repo so CI does not need to regenerate it.
Three detectable clauses: LIABILITY, CONFIDENTIALITY, PAYMENT.
"""
from docx import Document
import pathlib


def main():
    doc = Document()
    doc.add_heading("MASTER SERVICES AGREEMENT", level=1)
    doc.add_paragraph(
        'This agreement is entered into as of January 1, 2026, between Acme Corp '
        '("Provider") and Beta Inc ("Customer").'
    )
    doc.add_heading("1. LIABILITY", level=2)
    doc.add_paragraph(
        "Each party's total liability under this agreement shall not exceed USD 100,000. "
        "In no event shall either party be liable for any indirect, incidental, special, "
        "or consequential damages arising out of or related to this agreement."
    )
    doc.add_heading("2. CONFIDENTIALITY", level=2)
    doc.add_paragraph(
        "Each party shall hold the other party's confidential information in strict "
        "confidence for five (5) years following termination of this agreement. "
        "Confidential information shall not be disclosed to any third party without "
        "the prior written consent of the disclosing party."
    )
    doc.add_heading("3. PAYMENT", level=2)
    doc.add_paragraph(
        "Customer shall pay Provider within thirty (30) days of receipt of invoice. "
        "Late payments shall accrue interest at 1.5% per month or the maximum rate "
        "permitted by applicable law, whichever is lower."
    )
    doc.add_paragraph(
        "This agreement shall be governed by and construed in accordance with the "
        "laws of the Republic of Indonesia. Any disputes shall be resolved by the "
        "courts of Jakarta."
    )
    doc.add_paragraph(
        "This agreement constitutes the entire understanding between the parties "
        "with respect to the subject matter hereof and supersedes all prior agreements."
    )

    out = pathlib.Path(__file__).parent / "synth-contract.docx"
    doc.save(str(out))
    print(f"Wrote {out} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
