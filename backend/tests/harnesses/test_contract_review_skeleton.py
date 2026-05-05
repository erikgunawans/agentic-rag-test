"""Phase 22 / Plan 22-06 — Unit tests for Contract Review harness skeleton.

12 test cases (matching plan behavior block + ISSUE-09 + ISSUE-14):
 1. Flag gating — both flags True → harness registered as 'contract-review'
 2. Flag gating — either flag False → NOT registered (D-16 off-mode invariant)
 3. 9-phase shape — exact phase names in order (CR-01..08 + filter)
 4. Phase 1 type PROGRAMMATIC + executor is _phase1_intake
 5. Phase 2 type LLM_SINGLE + output_schema is ContractClassification
 6. CR-01 happy path — synthetic DOCX → extracts text, returns content + source_file
 7. CR-01 no upload → returns {"error": "no_uploaded_file", "code": "NO_UPLOAD"}
 8. ContractClassification rejects empty parties list (min_length=2 Pydantic)
 9. ContractClassification rejects single party (min_length=2)
10. ContractClassification accepts valid fully-populated instance
11. ISSUE-09: contract_review_enabled=True but tool_registry_enabled=False → RuntimeError
12. ISSUE-14: contract_review.py comment block mentions deploy-order constraint

Run:
    cd backend && source venv/bin/activate && \\
        pytest tests/harnesses/test_contract_review_skeleton.py -v --tb=short
"""
from __future__ import annotations

import io
import importlib
import sys
from unittest.mock import AsyncMock, MagicMock, patch

import pytest


# ---------------------------------------------------------------------------
# Fixtures / helpers
# ---------------------------------------------------------------------------

@pytest.fixture(autouse=True)
def _reset_registry():
    """Clear the harness registry before and after each test."""
    from app.services import harness_registry
    harness_registry._reset_for_tests()
    yield
    harness_registry._reset_for_tests()


def _reload_contract_review(
    harness_enabled: bool = True,
    contract_review_enabled: bool = True,
    tool_registry_enabled: bool = True,
):
    """Reload app.harnesses.contract_review under a patched settings object."""
    mock_settings = MagicMock()
    mock_settings.harness_enabled = harness_enabled
    mock_settings.contract_review_enabled = contract_review_enabled
    mock_settings.tool_registry_enabled = tool_registry_enabled

    # Remove cached module so we get a fresh import
    mod_name = "app.harnesses.contract_review"
    sys.modules.pop(mod_name, None)

    with patch("app.config.get_settings", return_value=mock_settings):
        mod = importlib.import_module(mod_name)

    return mod


# ---------------------------------------------------------------------------
# Tests 1-2: Flag gating
# ---------------------------------------------------------------------------

class TestFlagGating:
    def test_both_flags_true_registers_harness(self):
        """Test 1: harness_enabled=True + contract_review_enabled=True → registered."""
        from app.services import harness_registry
        _reload_contract_review(harness_enabled=True, contract_review_enabled=True, tool_registry_enabled=True)
        assert harness_registry.get_harness("contract-review") is not None

    @pytest.mark.parametrize("harness_enabled,cr_enabled", [
        (False, True),
        (True, False),
        (False, False),
    ])
    def test_either_flag_false_does_not_register(self, harness_enabled, cr_enabled):
        """Test 2: D-16 off-mode — either flag False → NOT registered."""
        from app.services import harness_registry
        # When contract_review_enabled=False with harness_enabled=True,
        # no RuntimeError — it just doesn't register.
        # When harness_enabled=False, contract_review_enabled check is skipped.
        _reload_contract_review(
            harness_enabled=harness_enabled,
            contract_review_enabled=cr_enabled,
            tool_registry_enabled=True,
        )
        assert harness_registry.get_harness("contract-review") is None


# ---------------------------------------------------------------------------
# Tests 3-5: Definition shape
# ---------------------------------------------------------------------------

class TestDefinitionShape:
    def test_nine_phases_in_order(self):
        """Test 3: exactly 9 phases with correct names."""
        mod = _reload_contract_review()
        names = [p.name for p in mod.CONTRACT_REVIEW.phases]
        expected = [
            "intake",
            "classify",
            "gather-context",
            "load-playbook",
            "extract-clauses",
            "risk-analysis",
            "filter-redline-candidates",
            "redline-generation",
            "executive-summary",
        ]
        assert names == expected

    def test_phase1_type_and_executor(self):
        """Test 4: Phase 1 is PROGRAMMATIC with executor=_phase1_intake."""
        from app.harnesses.types import PhaseType
        mod = _reload_contract_review()
        phase1 = mod.CONTRACT_REVIEW.phases[0]
        assert phase1.phase_type == PhaseType.PROGRAMMATIC
        assert phase1.executor is mod._phase1_intake

    def test_phase2_type_and_schema(self):
        """Test 5: Phase 2 is LLM_SINGLE with output_schema=ContractClassification."""
        from app.harnesses.types import PhaseType
        mod = _reload_contract_review()
        phase2 = mod.CONTRACT_REVIEW.phases[1]
        assert phase2.phase_type == PhaseType.LLM_SINGLE
        assert phase2.output_schema is mod.ContractClassification


# ---------------------------------------------------------------------------
# Test 6: CR-01 happy path — synthetic DOCX
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_phase1_intake_extracts_docx_text():
    """Test 6: CR-01 with a synthetic DOCX returns content + source_file."""
    from docx import Document
    buf = io.BytesIO()
    d = Document()
    d.add_paragraph("This is a test contract between Acme Corp and Beta Inc.")
    d.add_paragraph("Effective date: 2026-01-01.")
    d.save(buf)
    synthetic = buf.getvalue()

    ws_mock = MagicMock()
    ws_mock.list_files = AsyncMock(return_value=[
        {
            "file_path": "test.docx",
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "source": "upload",
            "size_bytes": len(synthetic),
        },
    ])
    ws_mock.read_binary_file = AsyncMock(return_value=synthetic)

    mod = _reload_contract_review()

    with patch("app.harnesses.contract_review.WorkspaceService", return_value=ws_mock):
        result = await mod._phase1_intake(
            inputs={}, token="tok", thread_id="thr", harness_run_id="run"
        )

    assert "content" in result, f"Expected 'content' in result, got: {result}"
    assert "Acme Corp" in result["content"]
    assert result.get("source_file") == "test.docx"
    assert result.get("page_count", 0) >= 1


# ---------------------------------------------------------------------------
# Test 7: CR-01 no upload → error dict
# ---------------------------------------------------------------------------

@pytest.mark.asyncio
async def test_phase1_intake_no_upload_returns_error():
    """Test 7: CR-01 with no uploaded files returns NO_UPLOAD error dict."""
    ws_mock = MagicMock()
    ws_mock.list_files = AsyncMock(return_value=[])  # no uploads

    mod = _reload_contract_review()

    with patch("app.harnesses.contract_review.WorkspaceService", return_value=ws_mock):
        result = await mod._phase1_intake(
            inputs={}, token="tok", thread_id="thr", harness_run_id="run"
        )

    assert result.get("error") == "no_uploaded_file"
    assert result.get("code") == "NO_UPLOAD"


# ---------------------------------------------------------------------------
# Tests 8-10: ContractClassification Pydantic schema validation
# ---------------------------------------------------------------------------

class TestContractClassificationSchema:
    def _get_schema(self):
        mod = _reload_contract_review()
        return mod.ContractClassification

    def test_rejects_empty_parties_list(self):
        """Test 8: parties=[] raises ValidationError (min_length=2)."""
        ContractClassification = self._get_schema()
        with pytest.raises(Exception):
            ContractClassification(
                contract_type="MSA",
                parties=[],
                governing_law="Republic of Indonesia",
                jurisdiction="courts of Jakarta",
                summary="Master Services Agreement between no parties and no one.",
            )

    def test_rejects_single_party(self):
        """Test 9: parties=["Acme"] raises ValidationError (min_length=2)."""
        ContractClassification = self._get_schema()
        with pytest.raises(Exception):
            ContractClassification(
                contract_type="MSA",
                parties=["Acme Corp"],
                governing_law="Republic of Indonesia",
                jurisdiction="courts of Jakarta",
                summary="Master Services Agreement between Acme Corp and no one.",
            )

    def test_accepts_valid_instance(self):
        """Test 10: fully-populated valid instance constructs without error."""
        ContractClassification = self._get_schema()
        obj = ContractClassification(
            contract_type="NDA",
            parties=["Acme Corp", "Beta Inc"],
            effective_date="2026-01-01",
            expiration_date="2027-01-01",
            governing_law="Republic of Indonesia",
            jurisdiction="courts of Jakarta",
            summary="Non-disclosure agreement between Acme Corp and Beta Inc governing confidential information.",
        )
        assert obj.contract_type == "NDA"
        assert len(obj.parties) == 2
        assert obj.governing_law == "Republic of Indonesia"


# ---------------------------------------------------------------------------
# Test 11: ISSUE-09 — tool_registry_enabled=False + contract_review_enabled=True
#          → RuntimeError at import time
# ---------------------------------------------------------------------------

def test_issue09_refuses_register_without_tool_registry():
    """Test 11: ISSUE-09 — RuntimeError when contract_review_enabled=True but tool_registry_enabled=False."""
    from app.services import harness_registry
    harness_registry._reset_for_tests()

    with pytest.raises(RuntimeError, match="ISSUE-09"):
        _reload_contract_review(
            harness_enabled=True,
            contract_review_enabled=True,
            tool_registry_enabled=False,
        )


# ---------------------------------------------------------------------------
# Test 12: ISSUE-14 deploy-order constraint documented in module
# ---------------------------------------------------------------------------

def test_issue14_deploy_order_documented():
    """Test 12: contract_review.py contains deploy-order/ISSUE-14 comment."""
    import inspect
    mod = _reload_contract_review()
    # Check docstring mentions the deploy constraint
    source = inspect.getsource(mod)
    assert "ISSUE-14" in source or "deploy-order" in source, (
        "contract_review.py must mention ISSUE-14 or deploy-order constraint in a comment"
    )
