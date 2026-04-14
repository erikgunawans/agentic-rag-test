"""
Document Upload / Ingestion API tests.
Covers: DOC-01 through DOC-06, SEC-05, DEDUP-01 through DEDUP-05, FMT-01 through FMT-08
"""
import json
import os
import time
import uuid
import pytest

FIXTURES = os.path.join(os.path.dirname(__file__), "../fixtures")


def upload_txt(client, filename="sample.txt"):
    """Upload a text file with a unique suffix to avoid dedup collisions across tests."""
    path = os.path.join(FIXTURES, filename)
    with open(path, "rb") as f:
        content = f.read() + f"\n# {uuid.uuid4()}".encode()
    resp = client.post(
        "/documents/upload",
        files={"file": (filename, content, "text/plain")},
    )
    return resp


class TestDocumentUpload:
    """DOC-01: Upload returns 202 with pending status."""

    def test_upload_txt_returns_202(self, authed_client):
        resp = upload_txt(authed_client)
        assert resp.status_code == 202

    def test_upload_returns_required_fields(self, authed_client):
        resp = upload_txt(authed_client)
        body = resp.json()
        assert "id" in body
        assert "filename" in body
        assert body["status"] == "pending"

    def test_upload_unsupported_type_returns_400(self, authed_client):
        """DOC-03: Unsupported file type."""
        resp = authed_client.post(
            "/documents/upload",
            files={"file": ("archive.zip", b"fake content", "application/zip")},
        )
        assert resp.status_code == 400

    def test_upload_empty_file_returns_400(self, authed_client):
        """DOC-04: Empty file."""
        resp = authed_client.post(
            "/documents/upload",
            files={"file": ("empty.txt", b"", "text/plain")},
        )
        assert resp.status_code == 400

    def test_upload_without_auth_returns_401(self, client):
        """SEC-05."""
        with open(os.path.join(FIXTURES, "sample.txt"), "rb") as f:
            resp = client.post(
                "/documents/upload",
                files={"file": ("sample.txt", f, "text/plain")},
            )
        assert resp.status_code in (401, 403)


class TestDocumentIngestion:
    """DOC-02: Ingestion completes with chunks."""

    def test_ingestion_completes(self, authed_client):
        resp = upload_txt(authed_client)
        doc_id = resp.json()["id"]

        # Poll for completion (max 45s)
        for _ in range(45):
            docs = authed_client.get("/documents").json()
            doc = next((d for d in docs if d["id"] == doc_id), None)
            assert doc is not None, "Document disappeared from list"
            if doc["status"] in ("completed", "failed"):
                break
            time.sleep(1)

        assert doc["status"] == "completed", f"Expected completed, got: {doc['status']} — error: {doc.get('error_msg')}"
        assert (doc["chunk_count"] or 0) >= 1, "Expected at least 1 chunk"


class TestDocumentList:
    """DOC-05: GET /documents returns user's own documents."""

    def test_list_returns_200(self, authed_client):
        resp = authed_client.get("/documents")
        assert resp.status_code == 200

    def test_list_returns_array(self, authed_client):
        assert isinstance(authed_client.get("/documents").json(), list)

    def test_list_has_required_fields(self, authed_client):
        upload_txt(authed_client)
        docs = authed_client.get("/documents").json()
        assert len(docs) > 0
        doc = docs[0]
        for field in ("id", "filename", "file_size", "mime_type", "status", "created_at"):
            assert field in doc, f"Missing field: {field}"


class TestDocumentDelete:
    """DOC-06: DELETE removes document and chunks."""

    def test_delete_returns_204(self, authed_client):
        resp = upload_txt(authed_client)
        doc_id = resp.json()["id"]
        del_resp = authed_client.delete(f"/documents/{doc_id}")
        assert del_resp.status_code == 204

    def test_deleted_document_not_in_list(self, authed_client):
        doc_id = upload_txt(authed_client).json()["id"]
        authed_client.delete(f"/documents/{doc_id}")
        ids = [d["id"] for d in authed_client.get("/documents").json()]
        assert doc_id not in ids

    def test_delete_nonexistent_returns_404(self, authed_client):
        resp = authed_client.delete("/documents/00000000-0000-0000-0000-000000000000")
        assert resp.status_code == 404


class TestDocumentDedup:
    """Module 3: Record Manager — content hash deduplication (DEDUP-01 through DEDUP-05)."""

    def _upload_content(self, client, content: bytes, filename: str = "dedup-test.txt"):
        return client.post(
            "/documents/upload",
            files={"file": (filename, content, "text/plain")},
        )

    def _wait_completed(self, client, doc_id: str, timeout: int = 45):
        """Poll until document reaches completed or failed status."""
        for _ in range(timeout):
            docs = client.get("/documents").json()
            doc = next((d for d in docs if d["id"] == doc_id), None)
            if doc and doc["status"] in ("completed", "failed"):
                return doc
            time.sleep(1)
        return None

    def test_new_upload_has_duplicate_false(self, authed_client):
        """DEDUP-01: Fresh upload response includes duplicate: false."""
        content = f"unique dedup content {uuid.uuid4()}".encode()
        resp = self._upload_content(authed_client, content)
        assert resp.status_code == 202
        assert resp.json()["duplicate"] is False

    def test_duplicate_upload_returns_200(self, authed_client):
        """DEDUP-02: Re-uploading completed doc returns 200 with duplicate: true and same ID."""
        content = f"unique dedup content {uuid.uuid4()}".encode()
        resp1 = self._upload_content(authed_client, content)
        assert resp1.status_code == 202
        doc_id = resp1.json()["id"]

        # Wait for ingestion to complete before re-uploading
        doc = self._wait_completed(authed_client, doc_id)
        assert doc is not None and doc["status"] == "completed"

        resp2 = self._upload_content(authed_client, content)
        assert resp2.status_code == 200
        body = resp2.json()
        assert body["duplicate"] is True
        assert body["id"] == doc_id

    def test_different_content_same_name_returns_202(self, authed_client):
        """DEDUP-03: Same filename but different content is treated as a new document."""
        content_a = f"content A {uuid.uuid4()}".encode()
        content_b = f"content B {uuid.uuid4()}".encode()
        resp1 = self._upload_content(authed_client, content_a, "same-name.txt")
        assert resp1.status_code == 202
        resp2 = self._upload_content(authed_client, content_b, "same-name.txt")
        assert resp2.status_code == 202
        assert resp1.json()["id"] != resp2.json()["id"]

    def test_content_hash_in_list(self, authed_client):
        """DEDUP-04: GET /documents includes content_hash as 64-char hex string."""
        content = f"hash field test {uuid.uuid4()}".encode()
        resp = self._upload_content(authed_client, content)
        assert resp.status_code == 202
        doc_id = resp.json()["id"]

        docs = authed_client.get("/documents").json()
        doc = next((d for d in docs if d["id"] == doc_id), None)
        assert doc is not None
        assert "content_hash" in doc
        assert doc["content_hash"] is not None
        assert len(doc["content_hash"]) == 64

    def test_reupload_after_delete_returns_202(self, authed_client):
        """DEDUP-05: Re-uploading after deletion creates a new document."""
        content = f"reupload test {uuid.uuid4()}".encode()
        resp1 = self._upload_content(authed_client, content)
        assert resp1.status_code == 202
        doc_id1 = resp1.json()["id"]

        # Wait for completion then delete
        self._wait_completed(authed_client, doc_id1)
        authed_client.delete(f"/documents/{doc_id1}")

        # Re-upload same content — should be treated as new
        resp2 = self._upload_content(authed_client, content)
        assert resp2.status_code == 202
        assert resp2.json()["id"] != doc_id1


VALID_CATEGORIES = {"technical", "legal", "business", "academic", "personal", "other"}


class TestDocumentMetadata:
    """Module 4: Metadata Extraction (META-01 through META-06)."""

    def _upload_content(self, client, content: bytes, filename: str = "meta-test.txt"):
        return client.post(
            "/documents/upload",
            files={"file": (filename, content, "text/plain")},
        )

    def _wait_completed(self, client, doc_id: str, timeout: int = 60):
        """Poll until document reaches completed or failed status."""
        for _ in range(timeout):
            docs = client.get("/documents").json()
            doc = next((d for d in docs if d["id"] == doc_id), None)
            if doc and doc["status"] in ("completed", "failed"):
                return doc
            time.sleep(1)
        return None

    def test_completed_document_has_metadata(self, authed_client):
        """META-01: After ingestion completes, metadata field is a dict (not null)."""
        content = f"Introduction to Python Programming\n\nPython is a versatile programming language used for web development, data science, and automation. This guide covers the basics. {uuid.uuid4()}".encode()
        resp = self._upload_content(authed_client, content)
        assert resp.status_code == 202
        doc_id = resp.json()["id"]

        doc = self._wait_completed(authed_client, doc_id)
        assert doc is not None, "Document did not reach terminal status"
        assert doc["status"] == "completed", f"Expected completed, got: {doc['status']}"
        assert doc.get("metadata") is not None, "metadata field should be populated after completion"
        assert isinstance(doc["metadata"], dict), "metadata should be a dict"

    def test_metadata_has_required_fields(self, authed_client):
        """META-02: Metadata contains required keys."""
        content = f"Quarterly Business Report Q3 2024\n\nRevenue increased by 15% compared to last quarter. Key metrics show growth across all departments. {uuid.uuid4()}".encode()
        resp = self._upload_content(authed_client, content)
        assert resp.status_code == 202
        doc_id = resp.json()["id"]

        doc = self._wait_completed(authed_client, doc_id)
        assert doc is not None and doc["status"] == "completed"
        meta = doc["metadata"]
        assert meta is not None
        for field in ("title", "category", "tags", "summary"):
            assert field in meta, f"Missing metadata field: {field}"

    def test_metadata_tags_is_nonempty_list(self, authed_client):
        """META-03: tags is a list with at least 1 item."""
        content = f"Legal Contract Terms and Conditions\n\nThis agreement governs the use of services provided. All parties must comply with applicable law. {uuid.uuid4()}".encode()
        resp = self._upload_content(authed_client, content)
        assert resp.status_code == 202
        doc_id = resp.json()["id"]

        doc = self._wait_completed(authed_client, doc_id)
        assert doc is not None and doc["status"] == "completed"
        tags = doc["metadata"]["tags"]
        assert isinstance(tags, list), "tags must be a list"
        assert len(tags) >= 1, "tags must have at least 1 item"

    def test_metadata_category_is_valid(self, authed_client):
        """META-04: category is one of the allowed values."""
        content = f"Academic Research Paper: Transformer Neural Networks\n\nThis paper presents a novel approach to attention mechanisms in deep learning. {uuid.uuid4()}".encode()
        resp = self._upload_content(authed_client, content)
        assert resp.status_code == 202
        doc_id = resp.json()["id"]

        doc = self._wait_completed(authed_client, doc_id)
        assert doc is not None and doc["status"] == "completed"
        category = doc["metadata"]["category"]
        assert category in VALID_CATEGORIES, f"Invalid category: {category}"

    def test_get_metadata_endpoint(self, authed_client):
        """META-05: GET /documents/{id}/metadata returns the metadata object."""
        content = f"Personal Journal Entry\n\nToday was a great day. I learned about RAG systems and vector databases. {uuid.uuid4()}".encode()
        resp = self._upload_content(authed_client, content)
        assert resp.status_code == 202
        doc_id = resp.json()["id"]

        doc = self._wait_completed(authed_client, doc_id)
        assert doc is not None and doc["status"] == "completed"

        meta_resp = authed_client.get(f"/documents/{doc_id}/metadata")
        assert meta_resp.status_code == 200
        meta = meta_resp.json()
        assert meta is not None
        assert "title" in meta

    def test_get_metadata_nonexistent_returns_404(self, authed_client):
        """META-06: GET /documents/{nonexistent}/metadata returns 404."""
        resp = authed_client.get("/documents/00000000-0000-0000-0000-000000000000/metadata")
        assert resp.status_code == 404


# ── Module 5: Multi-Format Support ───────────────────────────────────────────

def upload_docx(client, filename="sample.docx"):
    """Generate a unique DOCX in-memory to avoid content-hash dedup collisions."""
    from docx import Document as DocxDocument
    import io
    doc = DocxDocument()
    doc.add_paragraph("RAG Test Document - DOCX Format")
    doc.add_paragraph("The capital of France is Paris.")
    doc.add_paragraph(f"Unique ID: {uuid.uuid4()}")
    buf = io.BytesIO()
    doc.save(buf)
    content = buf.getvalue()
    resp = client.post(
        "/documents/upload",
        files={"file": (filename, content,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    return resp


def upload_csv(client, filename="sample.csv"):
    """Upload a CSV file with a unique suffix to avoid dedup collisions."""
    path = os.path.join(FIXTURES, filename)
    with open(path, "rb") as f:
        content = f.read() + f"\n# {uuid.uuid4()}".encode()
    resp = client.post(
        "/documents/upload",
        files={"file": (filename, content, "text/csv")},
    )
    return resp


def upload_html(client, filename="sample.html"):
    """Upload an HTML file with a unique comment to avoid dedup collisions."""
    path = os.path.join(FIXTURES, filename)
    with open(path, "rb") as f:
        content = f.read() + f"\n<!-- {uuid.uuid4()} -->".encode()
    resp = client.post(
        "/documents/upload",
        files={"file": (filename, content, "text/html")},
    )
    return resp


def upload_json(client, filename="sample.json"):
    """Upload a JSON file with a unique field to avoid dedup collisions."""
    path = os.path.join(FIXTURES, filename)
    with open(path, "rb") as f:
        data = json.loads(f.read())
    data["_dedup"] = str(uuid.uuid4())
    content = json.dumps(data).encode()
    resp = client.post(
        "/documents/upload",
        files={"file": (filename, content, "application/json")},
    )
    return resp


class TestMultiFormatUpload:
    """Module 5: Multi-Format Support (FMT-01 through FMT-08)."""

    def _wait_completed(self, client, doc_id: str, timeout: int = 45):
        for _ in range(timeout):
            docs = client.get("/documents").json()
            doc = next((d for d in docs if d["id"] == doc_id), None)
            if doc and doc["status"] in ("completed", "failed"):
                return doc
            time.sleep(1)
        return None

    def test_docx_upload_accepted(self, authed_client):
        """FMT-01: DOCX upload returns 202."""
        resp = upload_docx(authed_client)
        assert resp.status_code == 202

    def test_docx_ingestion_completes(self, authed_client):
        """FMT-02: DOCX ingestion completes with chunks."""
        resp = upload_docx(authed_client)
        doc_id = resp.json()["id"]
        doc = self._wait_completed(authed_client, doc_id)
        assert doc is not None, "Document did not reach terminal status"
        assert doc["status"] == "completed", f"Got: {doc['status']} — {doc.get('error_msg')}"
        assert (doc["chunk_count"] or 0) >= 1

    def test_csv_upload_accepted(self, authed_client):
        """FMT-03: CSV upload returns 202."""
        resp = upload_csv(authed_client)
        assert resp.status_code == 202

    def test_csv_ingestion_completes(self, authed_client):
        """FMT-04: CSV ingestion completes with chunks."""
        resp = upload_csv(authed_client)
        doc_id = resp.json()["id"]
        doc = self._wait_completed(authed_client, doc_id)
        assert doc is not None
        assert doc["status"] == "completed", f"Got: {doc['status']} — {doc.get('error_msg')}"
        assert (doc["chunk_count"] or 0) >= 1

    def test_html_upload_accepted(self, authed_client):
        """FMT-05: HTML upload returns 202."""
        resp = upload_html(authed_client)
        assert resp.status_code == 202

    def test_html_ingestion_completes(self, authed_client):
        """FMT-06: HTML ingestion completes with chunks."""
        resp = upload_html(authed_client)
        doc_id = resp.json()["id"]
        doc = self._wait_completed(authed_client, doc_id)
        assert doc is not None
        assert doc["status"] == "completed", f"Got: {doc['status']} — {doc.get('error_msg')}"
        assert (doc["chunk_count"] or 0) >= 1

    def test_json_upload_accepted(self, authed_client):
        """FMT-07: JSON upload returns 202."""
        resp = upload_json(authed_client)
        assert resp.status_code == 202

    def test_json_ingestion_completes(self, authed_client):
        """FMT-08: JSON ingestion completes with chunks."""
        resp = upload_json(authed_client)
        doc_id = resp.json()["id"]
        doc = self._wait_completed(authed_client, doc_id)
        assert doc is not None
        assert doc["status"] == "completed", f"Got: {doc['status']} — {doc.get('error_msg')}"
        assert (doc["chunk_count"] or 0) >= 1
